import requests
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
import re
import os
from docx import Document
from urllib.parse import quote, urljoin

def rastreador_unificado():
    directorio = os.getcwd()
    hoy = datetime.utcnow() + timedelta(hours=2) 
    fecha_hoy_str = hoy.strftime("%d_%m_%Y")
    
    # Archivos de salida
    nombre_word_tic = os.path.join(directorio, f"Oposiciones_{fecha_hoy_str}.docx")
    nombre_word_dipu = os.path.join(directorio, f"Diputacion_Coruna_{fecha_hoy_str}.docx")
    api_key_proxy = os.environ.get("SCRAPER_API_KEY")
    
    # --- FILTROS TIC ---
    terminos_it = [r"\binformática\b", r"\binformático\b", r"\bprogramador\b", r"\bsoftware\b", 
                   r"\btic\b", r"\bsistemas de información\b", r"\bdixital\b", r"\bdigital\b", r"\bredes\b",
                   r"\btecnoloxía da información\b", r"\btecnología de la información\b"]
    terminos_genericos = [r"\bcuerpos y escalas\b", r"\bcorpos e escalas\b", r"\boferta de empleo público\b", 
                          r"\boferta de emprego público\b", r"\boep\b", r"\bcorpo superior\b", r"\bcuerpo superior\b",
                          r"\bescala de sistemas\b"]
    accion = ["convoca", "proceso selectivo", "oposición", "libre", "quenda", "prazas", "ingreso", "ferrol",
              "estabilización", "oferta de empleo", "oep", "oferta de emprego", "personal laboral", "funcionario"]
    
    # --- FILTROS DIPUTACIÓN ---
    entidad_cast = "diputación provincial de a coruña"
    entidad_gal = "deputación da coruña"
    rrhh_terminos = ["recursos humanos", "rrhh", "recursos humans", "oferta de empleo", "oferta de emprego", "proceso selectivo"]

    # Inicialización de documentos
    doc_tic = Document()
    doc_tic.add_heading(f'Boletín TIC - {hoy.strftime("%d/%m/%Y")}', 0)
    anuncios_tic_directos = {} 
    anuncios_posibles = {}

    doc_dipu = Document()
    doc_dipu.add_heading(f'Alertas Diputación A Coruña - {hoy.strftime("%d/%m/%Y")}', 0)
    anuncios_diputacion = {}

    total_llamadas_proxy = 0
    print(f"\n--- 🛰️ INICIANDO BÚSQUEDA UNIFICADA (TIC + DIPUTACIÓN) ---")
    
    sesion = requests.Session()
    sesion.headers.update({'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36'})

    for i in range(7):
        fecha = hoy - timedelta(days=i)
        f_str = fecha.strftime("%d/%m/%Y")
        print(f"🔍 Analizando {f_str}...")
        
        dia_semana = fecha.weekday() 
        urls = {}
        
        if dia_semana != 6: 
            urls["BOE"] = fecha.strftime("https://www.boe.es/boe/dias/%Y/%m/%d/")
        if dia_semana not in [5, 6]: 
            urls["BOP Coruña"] = f"https://bop.dacoruna.gal/bopportal/cambioBoletin.do?fechaInput={f_str}"
        
        for sec in range(1, 7):
            urls[f"DOG Sec {sec}"] = f"https://www.xunta.gal/diario-oficial-galicia/mostrarContenido.do?ruta=/{fecha.year}/{fecha.strftime('%Y%m%d')}/Secciones{sec}_gl.html"
        
        for fuente, url in urls.items():
            try:
                if (fuente.startswith("DOG") or fuente == "BOE") and api_key_proxy:
                    total_llamadas_proxy += 1
                    res = requests.get(f"http://api.scraperapi.com?api_key={api_key_proxy}&url={quote(url, safe='')}&render=false", timeout=45) 
                else:
                    res = sesion.get(url, timeout=20)
                
                if res.status_code != 200: continue
                    
                sopa = BeautifulSoup(res.text, 'html.parser')
                elementos_analizar = []
                
                # --- EXTRACCIÓN ROBUSTA (Motor TIC) ---
                if "BOP Coruña" in fuente:
                    for item in sopa.find_all(['li', 'p']):
                        link = item.find('a', href=True)
                        if link:
                            id_match = re.search(r'(\d{4}_\d+)', link['href'])
                            id_archivo = id_match.group(1) if id_match else "0000_0000"
                            u = f"https://bop.dacoruna.gal/bopportal/publicado/{fecha.strftime('%Y/%m/%d')}/{id_archivo}.pdf"
                            
                            texto_item = item.get_text(separator=" ")
                            municipio = "BOP Coruña"
                            prev_node = item.find_previous(['h1', 'h2', 'h3', 'p', 'div'])
                            if prev_node: municipio = prev_node.get_text(strip=True)
                            
                            elementos_analizar.append((f"🏢 {municipio} | {texto_item}", u))
                else:
                    for item in sopa.find_all(['li', 'p']):
                        link = item.find('a', href=True)
                        if link:
                            elementos_analizar.append((item.get_text(separator=" "), urljoin(url, link['href'])))

                # --- LÓGICA BIFURCADA DE CLASIFICACIÓN ---
                for texto, url_final in elementos_analizar:
                    texto_limpio = texto.strip()
                    if len(texto_limpio) < 15: continue 
                    txt_min = texto_limpio.lower()
                    
                    huella = re.sub(r'\W+', '', texto_limpio)[:200]
                    datos = {'texto': texto_limpio, 'fuente': fuente, 'fecha': f_str, 'url': url_final}

                    # EVALUACIÓN 1: Puestos TIC
                    tiene_it = any(re.search(t, txt_min) for t in terminos_it)
                    tiene_generico = any(re.search(g, txt_min) for g in terminos_genericos)
                    tiene_accion = any(a in txt_min for a in accion)
                    
                    if (tiene_it or tiene_generico) and tiene_accion:
                        if tiene_it: 
                            anuncios_tic_directos[huella] = datos
                        else: 
                            anuncios_posibles[huella] = datos

                    # EVALUACIÓN 2: Diputación
                    tiene_entidad = (entidad_cast in txt_min or entidad_gal in txt_min)
                    if tiene_entidad:
                        is_boe = (fuente == "BOE")
                        tiene_rrhh = any(r in txt_min for r in rrhh_terminos)
                        # Aplica doble validación si no es BOE
                        if is_boe or tiene_rrhh:
                            anuncios_diputacion[huella] = datos

            except Exception: 
                continue

    # --- ESCRITURA DOC 1: TIC ---
    doc_tic.add_heading('🎯 Búsqueda Directa (Puestos TIC)', level=1)
    for d in anuncios_tic_directos.values():
        doc_tic.add_paragraph(f"{d['fuente']} - {d['fecha']}\n{d['texto']}\n🔗 {d['url']}\n" + "-"*30)
    
    doc_tic.add_heading('⚠️ Posibles Oposiciones', level=1)
    for d in anuncios_posibles.values():
        doc_tic.add_paragraph(f"{d['fuente']} - {d['fecha']}\n{d['texto']}\n🔗 {d['url']}\n" + "-"*30)
    doc_tic.save(nombre_word_tic)

    # --- ESCRITURA DOC 2: DIPUTACIÓN ---
    for d in anuncios_diputacion.values():
        p = doc_dipu.add_paragraph()
        p.add_run(f"📌 {d['fuente']} - {d['fecha']}").bold = True
        doc_dipu.add_paragraph(d['texto'])
        doc_dipu.add_paragraph(f"🔗 {d['url']}")
        doc_dipu.add_paragraph("-" * 30)
    
    if not anuncios_diputacion:
        doc_dipu.add_paragraph("\nℹ️ Sin novedades con estos filtros en los últimos 7 días.")
    doc_dipu.save(nombre_word_dipu)

    print(f"\n✅ ¡Hecho! Resultados TIC: {len(anuncios_tic_directos)} directos, {len(anuncios_posibles)} posibles.")
    print(f"✅ ¡Hecho! Resultados Diputación: {len(anuncios_diputacion)} anuncios.")
    print(f"📡 (Llamadas Proxy: {total_llamadas_proxy})")

if __name__ == "__main__":
    rastreador_unificado()
